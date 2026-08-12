"""Generate the stakeholder tenancy explainer Word document."""
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)
LIGHT = RGBColor(0xED, 0xF2, 0xF8)


def set_base_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, space_before in (
        ("Heading 1", 18, ACCENT, 14),
        ("Heading 2", 14, ACCENT, 12),
        ("Heading 3", 12, ACCENT, 8),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(4)


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Multi-Tenancy in the SITA Security Platform")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("How data isolation and delegated access are handled across the national and provincial estate")
    run.font.size = Pt(14)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Prepared for stakeholders · {date.today().strftime('%d %B %Y')}")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    doc.add_page_break()


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items, bold_lead=True):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_lead and " — " in item:
            lead, rest = item.split(" — ", 1)
            p.add_run(lead + " — ").bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def role_table(doc, rows, header):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def callout(doc, label, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Light Shading Accent 1"
    cell = t.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(text)


def main():
    doc = Document()
    set_base_styles(doc)
    sec = doc.sections[0]
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    cover(doc)

    # ------------------------------------------------------------------ #
    heading(doc, "1. Purpose of this document", 1)
    para(
        doc,
        "This document explains, in non-technical language, how the SITA Security "
        "Platform keeps data isolated between tenants and how access is delegated. "
        "It is intended for executives, departmental security officers, provincial "
        "coordinators, and governance stakeholders who need to understand the "
        "safeguards in place without reading source code."
    )
    para(
        doc,
        "A \u201ctenant\u201d is a government department (an administrative organ of state) "
        "served by SITA \u2014 for example National Treasury, the Department of Home "
        "Affairs, or any of the 113 provincial departments. Every tenant only ever "
        "sees its own data unless a user is explicitly granted a wider scope."
    )

    # ------------------------------------------------------------------ #
    heading(doc, "2. Why multi-tenancy matters", 1)
    para(
        doc,
        "SITA consolidates public-sector ICT under a legislative mandate that makes "
        "every national and provincial department a mandatory client. The platform "
        "must therefore serve many customers at once, and the central concern is "
        "isolation: "
    )
    bullets(doc, [
        "Confidentiality — \u2014 A department must never see another department's security findings, risk scores, or alerts.",
        "Least privilege — \u2014 Each user sees exactly the portion of the estate their role requires, and nothing more.",
        "Delegated administration — \u2014 Each admin manages only the users and scope beneath their own node in the hierarchy.",
        "Assurance — \u2014 Every action is attributable to an actor and bound to a tenant scope for audit and, where needed, external oversight.",
    ])

    # ------------------------------------------------------------------ #
    heading(doc, "3. The tenancy model", 1)
    heading(doc, "3.1 Three-tier hierarchy", 2)
    para(
        doc,
        "Tenants are organised as a strict tree, mirroring SITA's own customer "
        "structure and its Managed Services hierarchy:"
    )
    bullets(doc, [
        "System identity (Tenant) — \u2014 the national department or provincial administration that holds the account.",
        "Sub-Tenant (Department) — \u2014 the department itself, e.g. Department of Home Affairs or Gauteng Department of Health.",
        "Asset (Application / Database) — \u2014 the specific applications and databases belonging to that department.",
    ])
    para(
        doc,
        "National departments also carry a further organisational level, Branches "
        "(the Deputy Director-General-led units inside a department, e.g. the DHA "
        "Information Services / CIO branch). Provincial departments are scoped at "
        "the province level and do not carry a branch subtree."
    )

    heading(doc, "3.2 Coverage of the estate", 2)
    para(
        doc,
        "The tenancy catalog models the full SITA customer estate: 43 national "
        "departments and 113 provincial departments across all nine provinces "
        "(156 public clients), each with its own derived application and database "
        "assets. Department identifiers are stable text keys (slugs), so tenancy "
        "resolution never depends on database lookups and scoping stays fast and "
        "deterministic."
    )

    heading(doc, "3.3 How data is labelled", 2)
    para(
        doc,
        "Every warehouse record that belongs to a tenant carries its owning "
        "department and branch as explicit data columns, written at ingestion time. "
        "This means a user's access can be enforced with a simple filter on each "
        "query \u2014 no expensive cross-table joins are needed to determine who may see a "
        "row. Provincial departments additionally resolve to a province, so "
        "assigning a scope to a whole province automatically expands to every "
        "department inside it."
    )

    # ------------------------------------------------------------------ #
    heading(doc, "4. Who can see what (roles and scopes)", 1)
    para(
        doc,
        "Access is granted through roles, and each role is scoped to a set of "
        "departments, branches, or provinces. Scoping is set by an administrator "
        "and stored with the user's identity, then enforced on every single read "
        "and write by the platform (not just hidden in the user interface)."
    )

    heading(doc, "4.1 Role families", 2)
    role_table(doc, [
        ("Nationwide", "exec, compliance, sre, admin, transversal-admin", "Whole-estate visibility and management (senior leadership, governance, service operations)."),
        ("Department", "soc, appsec, dbsec", "Operational security roles scoped to assigned departments (and optionally branches)."),
        ("Province", "province-soc-lead, province-dept-admin, local-appsec", "Provincial personas scoped to a province's full department set, e.g. a Gauteng SOC lead."),
        ("Admin tiers", "dept-admin, branch-admin, province-dept-admin", "Delegated user administration at a department, branch, or province node."),
    ], ["Family", "Roles", "What it means"])

    heading(doc, "4.2 The scoping rule", 2)
    callout(
        doc,
        "Rule",
        "A user's scope is their assigned departments plus any branches inside "
        "those departments. A province scope expands to every department in that "
        "province. If no scope is set, nationwide roles see the whole estate; "
        "department roles fail closed (they see nothing) until an administrator "
        "assigns a scope."
    )

    # ------------------------------------------------------------------ #
    heading(doc, "5. Delegated administration", 1)
    para(
        doc,
        "User management is delegated along the same tenancy tree, mirroring SITA's "
        "Managed Services hierarchy. Each administrator tier is a node in the tree "
        "and its authority is exactly the subtree beneath that node."
    )
    role_table(doc, [
        ("Tier 4 · system admin", "estate root", "National superadmin; grants anything across the whole estate."),
        ("Tier 3 · transversal-admin", "estate or assigned scope", "Grants department/branch admin tiers and operational roles across its scope."),
        ("Tier 2 · dept-admin / province-dept-admin", "department or province", "Grants operational department roles (and branch-admin at national level) within its departments/province."),
        ("Tier 1 · branch-admin", "branch", "Grants operational department roles within its branches."),
    ], ["Tier", "Scope", "Authority"])

    para(
        doc,
        "Delegation is strictly one-way down the tree. An administrator can never "
        "grant a role at or above their own level, and can never widen a scope "
        "beyond their own subtree. A department admin therefore can never reach "
        "another department's administration. These rules are enforced "
        "programmatically \u2014 they are not a matter of policy or convention."
    )

    # ------------------------------------------------------------------ #
    heading(doc, "6. How isolation is enforced", 1)
    para(
        doc,
        "Isolation is enforced at the data layer, on the server, so it holds "
        "regardless of which client, tool, or report is used:"
    )
    bullets(doc, [
        "Enforced on every query — \u2014 each data access adds a tenant filter derived from the caller's identity, so a provincial caller's SQL cannot reach rows outside its province.",
        "Scoping is additive — \u2014 user-provided parameters (e.g. filters in a risk report) can only narrow a scope, never widen it.",
        "No joins required — \u2014 tenant ownership is denormalised onto each row at write time, so scoping cannot be bypassed through a missing relationship.",
        "Blinded benchmarking — \u2014 peer comparisons (e.g. provincial benchmarks) return peers as anonymous \u201cPeer Province A / B\u201d rather than exposing another tenant's identity or data.",
        "Fail closed — \u2014 a department user with no assigned scope sees nothing; the safe default is no access, not full access.",
    ])

    # ------------------------------------------------------------------ #
    heading(doc, "7. Provincial governance", 1)
    para(
        doc,
        "SITA's provincial mandate is modelled explicitly. Each of the nine "
        "provinces (Gauteng, Western Cape, KwaZulu-Natal, etc.) owns a set of "
        "provincial departments \u2014 11 to 15 per province, totalling 113. "
        "Provincial roles are scoped by province, and a province scope expands "
        "automatically to the province's full department set. Provincial data "
        "rolls up into the same national ministry and cluster reporting "
        "structure, so whole-of-government reporting remains possible while "
        "day-to-day access stays provincial."
    )

    # ------------------------------------------------------------------ #
    heading(doc, "8. Audit and oversight", 1)
    para(
        doc,
        "Admin actions that change scope, roles, or user records are written to an "
        "immutable audit trail that binds each action to the acting user and the "
        "tenant scope affected, together with an integrity hash that makes "
        "tampering detectable. Secure export capabilities for oversight bodies "
        "(e.g. the Auditor-General) record the exporting actor, the scope of the "
        "export, and a verifiable payload hash, so every disclosure is attributable."
    )

    # ------------------------------------------------------------------ #
    heading(doc, "9. What this means for you", 1)
    role_table(doc, [
        ("Executive / Governance", "You can rely on whole-of-estate dashboards while individual departments remain isolated; oversight and reporting routes are available."),
        ("Departmental SOC / AppSec / DBSec", "You operate only within your department's scope; provincial SOC leads see all departments in their province."),
        ("Department / Provincial administrators", "You administer only the users and scope beneath your own node; you cannot reach peers or superiors."),
        ("Provincial coordinators", "Province-wide scope is assigned once and expands to all departments in that province automatically."),
        ("Audit / Compliance", "Every admin action and export is bound to an actor and a tenant scope with tamper-evident hashes."),
    ], ["Stakeholder", "What you get"])

    # ------------------------------------------------------------------ #
    heading(doc, "10. Frequently asked questions", 1)
    heading(doc, "Can a department ever see another department's data?", 3)
    para(doc, "No. Every query is scoped server-side to the caller's assigned departments, branches, or provinces. There is no interface or parameter that can widen that scope.")
    heading(doc, "How do provincial users see multiple departments without extra setup?", 3)
    para(doc, "A province-scoped role (e.g. province-soc-lead) is assigned once; the platform expands that province to all of its departments automatically.")
    heading(doc, "What happens if a user has no scope assigned?", 3)
    para(doc, "A nationwide role sees the whole estate; any department- or province-scoped role sees nothing until an administrator assigns a scope. The default is safe.")
    heading(doc, "Can a delegated admin grant themselves more access?", 3)
    para(doc, "No. An administrator can only grant roles at or below their own tier and only within their own subtree. The system refuses grants that exceed the caller's authority.")
    heading(doc, "Is the isolation only at the user-interface level?", 3)
    para(doc, "No. Isolation is enforced in the data-access layer on the server. Hiding buttons in the interface is only a convenience; the security boundary is the data itself.")

    heading(doc, "11. Document control", 1)
    role_table(doc, [
        ("Version", "1.0"),
        ("Status", "For stakeholder review"),
        ("Audience", "Executives, governance, department security, provincial coordinators"),
        ("Basis", "Platform tenancy model (national + provincial estate, role hierarchy, delegated administration)"),
    ], ["Field", "Value"])

    out = r"C:\sita-platform\docs\Tenancy_Stakeholder_Overview.docx"
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
