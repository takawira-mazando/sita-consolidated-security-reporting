"""Generate the full system documentation Word document."""
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)


def set_base_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (("Heading 1", 18, ACCENT), ("Heading 2", 14, ACCENT), ("Heading 3", 12, ACCENT)):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SITA Consolidated Security Reporting")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Complete System Documentation")
    r.font.size = Pt(15)
    r.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{date.today().strftime('%d %B %Y')} · Multi-tenancy release (a53f829)")
    r.font.color.rgb = GREY
    doc.add_page_break()


def h(doc, text, lvl=1):
    doc.add_heading(text, lvl)


def p(doc, text):
    doc.add_paragraph(text)


def b(doc, items):
    for it in items:
        para = doc.add_paragraph(style="List Bullet")
        if " — " in it:
            lead, rest = it.split(" — ", 1)
            para.add_run(lead + " — ").bold = True
            para.add_run(rest)
        else:
            para.add_run(it)


def table(doc, rows, header=None):
    t = doc.add_table(rows=1, cols=len(header) if header else len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        for i, hh in enumerate(header):
            t.rows[0].cells[i].paragraphs[0].add_run(hh).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


def main():
    doc = Document()
    set_base_styles(doc)
    sec = doc.sections[0]
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    cover(doc)

    h(doc, "1. System overview", 1)
    p(doc, "SITA Consolidated Security Reporting is a full-stack security intelligence platform that unifies application security, database security, SOC operations, compliance and executive reporting. It runs as a multi-tenant government platform on SITA's mandate: 43 national departments and 113 provincial departments across 9 provinces (156 public clients).")
    b(doc, [
        "Ingestion — Four OEM sources (HCL AppScan, Imperva DAM/WAF, API Security, compliance CSV) plus a deterministic synthetic feed until live access exists.",
        "Warehouse — PostgreSQL 16 with staging, warehouse, archive, audit and identity schemas.",
        "Engines — Alert, risk (4-signal fusion 0-100) and compliance (POPIA + ISO 27001) engines.",
        "Dispatch — Email (Jinja2), Microsoft Teams Adaptive Cards, PagerDuty Events v2.",
        "Reporting — FastAPI REST API + React SPA with six role-gated dashboards.",
        "Tenancy — Province → Department → Branch → App isolation with delegated admin and tenant-safe analytics.",
    ])
    table(doc, [
        ("Frontend", "React 18, TypeScript, Vite, Recharts, Axios"),
        ("Backend", "Python 3.12, FastAPI, SQLAlchemy 2 async, Pydantic v2, pandas"),
        ("Data", "PostgreSQL 16 (asyncpg), Redis 7"),
        ("Auth", "JWT (HS256 dev / RS256 Auth0)"),
        ("Monitoring", "Prometheus, Grafana, structlog"),
        ("Infra", "Docker Compose, nginx, GHCR, GitHub Actions"),
    ], ["Layer", "Technology"])

    h(doc, "2. Architecture", 1)
    p(doc, "The system runs as eight Docker Compose services: postgres, redis, backend (FastAPI on :8000), ingestion, processing, dispatch, analytics, and frontend (nginx on :3000). The end-to-end flow is: OEM sources → staging.raw_records → normaliser → warehouse.findings → alert/risk/compliance engines → warehouse.* → REST API → React dashboards.")
    p(doc, "Repository layout: backend/ (app package, tests, requirements), frontend/ (React SPA), infrastructure/ (docker-compose, nginx, monitoring, scripts), docs/, .github/workflows/.")

    h(doc, "3. Multi-tenancy model", 1)
    p(doc, "A tenant is a government department. The tenant tree is Province → Department → Branch → App/Database. Every warehouse row denormalises department_id/branch_id at write time, so tenant_filter scoping needs no joins. Provincial departments are branchless; a province scope expands to the province's full department set.")
    h(doc, "3.1 Scoping roles", 2)
    table(doc, [
        ("Nationwide", "exec, compliance, sre, admin, transversal-admin", "Whole estate (empty scope = whole estate)"),
        ("Department", "soc, appsec, dbsec", "Scoped to department_ids / branch_ids"),
        ("Province", "province-soc-lead, province-dept-admin, local-appsec", "Scoped by province_ids; expands to province departments"),
    ], ["Family", "Roles", "Scope"])
    h(doc, "3.2 Delegated admin tiers", 2)
    table(doc, [
        ("Tier 4", "system admin", "Estate root — grants anything"),
        ("Tier 3", "transversal-admin", "Operational roles + dept/branch-admin across scope"),
        ("Tier 2", "dept-admin / province-dept-admin", "Operational roles (+ branch-admin nationally)"),
        ("Tier 1", "branch-admin", "Operational roles only"),
    ], ["Tier", "Role", "Authority"])
    p(doc, "Enforcement: tenant_filter(claims, model) adds the WHERE clause on every read in auth.py. Delegation is one-way down via can_manage; scope cannot exceed the caller's subtree. Fail-closed: scope-less department users see nothing.")
    p(doc, "Full stakeholder explainer: docs/Tenancy_Stakeholder_Overview.docx.")

    h(doc, "4. Authentication & authorization", 1)
    p(doc, "Users authenticate via POST /api/v1/auth/login (or demo-login with a tenancy-scope override) and receive a JWT. verify_token decodes it into JWTClaims; require_roles checks the role permission matrix. On startup the lifespan bootstraps the superadmin and seeds demo users.")
    table(doc, [
        ("admin", "*"), ("exec", "risks, compliance, alerts, alerts_read/write, findings, dashboard"),
        ("soc", "findings, alerts_read/write, dashboard"), ("appsec", "risks, findings"),
        ("dbsec", "risks, findings, alerts_read"), ("compliance", "compliance"),
        ("sre", "admin_read/write, alerts_read, dashboard"),
        ("transversal-admin", "admin_read/write"), ("dept-admin", "admin_read/write"),
        ("branch-admin", "admin_read/write"), ("province-soc-lead", "findings, alerts_read/write, dashboard"),
        ("province-dept-admin", "admin_read/write"), ("local-appsec", "risks, findings"),
    ], ["Role", "Permissions"])

    h(doc, "5. Data pipeline", 1)
    b(doc, [
        "Connectors — BaseConnector (authenticate → poll → parse → validate → run) with concrete AppScan, Imperva, API-Security and Compliance connectors; APScheduler, circuit breaker, token bucket, retry, DLQ.",
        "Synthetic feed — deterministic (seed=42) OEM-shaped payloads until live access; cutover is a source swap at the ingestion boundary.",
        "Normaliser — field_mapper.yaml JSONPath maps, severity normalisation, UTC coercion, dedup by source+external_id, DLQ routing.",
        "Alert engine — six rules (risk_critical, new_critical_cve, shadow_api_detected, compliance_drop, imperva_violation_spike, connector_auth_failure); Redis throttle suppression; enricher adds owner/team/priority.",
        "Risk engine — R = (0.35·Sa + 0.25·Si + 0.20·E + 0.20·C); buckets safe (<45) / monitored (45-69) / critical (>=70).",
        "Compliance — POPIA 8 conditions (security safeguards 0.30) + ISO 27001 A.5-A.16.",
        "Dispatch — email/Teams/PagerDuty adapters + async worker writing dispatch_log and last_dispatched_at.",
    ])

    h(doc, "6. Database schema", 1)
    p(doc, "Six schemas: staging, warehouse, archive, audit, identity. Key tables: identity.users (with province_ids/department_ids/branch_ids JSONB scope lists) and identity.persons; audit.action_audit (actor, action, tenant_scope, payload_hash); staging.batch_runs/raw_records/rejected_records; warehouse.provinces/departments/branches, findings, risk_scores, alerts, dispatch_log, compliance_snapshots/gaps, api_endpoints, waf_blocks, system_metrics, agents, slo_metrics, database_inventory, connector_health; archive copies of risk_scores/alerts/compliance_snapshots/findings. Full DDL in infrastructure/scripts/init-db.sql; interactive ER diagram in tree.html.")

    h(doc, "7. REST API reference", 1)
    p(doc, "Prefixes: /api/v1 (core), /admin (admin), /health. Interactive docs at /docs.")
    table(doc, [
        ("Auth", "POST /auth/login, POST /auth/demo-login, GET /auth/demo-accounts, GET /auth/tenancy, GET /auth/me"),
        ("Risks", "GET /risks, GET /risks/by-cluster, GET /risks/by-ministry, GET /risks/{app_name}/trend"),
        ("Findings", "GET /findings, GET /findings/{id}"),
        ("Compliance", "GET /compliance, /compliance/gaps, /compliance/evidence"),
        ("Alerts", "GET /alerts, GET /alerts/{id}/dispatch, PATCH /alerts/{id}/acknowledge, PATCH /alerts/{id}/resolve"),
        ("Dashboard", "GET /dashboard/summary"),
        ("Metrics", "appsec/waf, appsec/api-exposure, appsec/fix-rate, soc/slo, sre/system, sre/agents, dbsec/inventory, compliance/trend, compliance/calendar"),
        ("Exports", "GET /exports/findings.csv, /exports/risk-scores.csv, /exports/alerts.csv, POST /exports/ag-compliance, GET /exports/ag-compliance/verify/{hash}"),
        ("Benchmark", "GET /benchmark/province (anonymised peers)"),
        ("Admin", "GET /admin/departments, /branches, /connectors, /dead-letter, /users, /persons; POST /admin/connectors/{name}/reset, /dead-letter/reprocess/{id}, /users, /hr/sync; PATCH /admin/users/{id}; DELETE /admin/users/{id}"),
    ], ["Group", "Endpoints"])

    h(doc, "8. Frontend application", 1)
    p(doc, "React 18 + TypeScript + Vite SPA. Pages: Exec, Soc, AppSec, DbSec, Compliance, Sre dashboards, AdminUsersPage (user management with province/department/branch scope pickers), LoginPage, HomePage. Component kit in components/dashboard (StatCard, Panel, Gauge, MiniBarChart, BarRow, Chip, OemTag, DashHeader, ExplainOverlay). Typed API clients in src/api, per-dashboard hooks in src/hooks, role-aware navigation in RoleNav, auth state in AuthContext.")

    h(doc, "9. Infrastructure & deployment", 1)
    b(doc, [
        "Docker Compose — 8 services; backend image shared by backend/ingestion/processing/dispatch/analytics.",
        "nginx — HTTPS-only TLS 1.3, HSTS/CSP headers, rate limit on /api/v1/auth/ (5 r/s), proxies /api and /admin to backend, / to frontend.",
        "Scripts — init-db.sql (full schema), seed-demo.sql, maintenance.sql (archive TTL), backup.sh (pg_dump -Fc daily), deploy.sh.",
        "Backups & DR — daily dumps with retention; WAL/PITR path for RTO < 15 min.",
    ])

    h(doc, "10. Monitoring & observability", 1)
    p(doc, "Prometheus scrapes backend /metrics, Postgres (9187) and Redis (9121); Grafana dashboard sita-pipeline.json visualises the pipeline; structlog provides structured logging; /health reports uptime.")

    h(doc, "11. CI/CD", 1)
    p(doc, "ci.yml (push to main/develop, PR to main): lint (ruff + eslint), test (pytest against Postgres 16 + Redis 7, frontend build), build (docker). publish-ghcr.yml publishes backend/frontend images to GHCR on main.")

    h(doc, "12. Security model", 1)
    b(doc, [
        "Server-side tenant_filter isolation on every query, fail-closed defaults.",
        "One-way-down delegated admin; scope cannot exceed caller subtree.",
        "JWT HS256 dev / RS256 Auth0; config guards reject default dev secrets in non-dev.",
        "AG exports: PII stripped, aggregation-only, SHA-256 hash bound to actor + tenant scope (tamper-evident audit trail).",
        "Blinded provincial benchmarking — peers shown as anonymous aggregates.",
    ])

    h(doc, "13. Operations runbook", 1)
    table(doc, [
        ("Start stack", "docker compose -f infrastructure/docker-compose.yml up --build"),
        ("Backend local", "cd backend && python -m uvicorn app.main:app --port 8000"),
        ("Frontend local", "cd frontend && npm run dev -- --host 0.0.0.0 --port 3000"),
        ("Tenancy migration", "cd backend && python _migrate_tenancy.py && python _migrate_province.py"),
        ("Seed simulated data", "cd backend && python -m app.entrypoints.seed_simulated"),
        ("Health check", "curl http://localhost:8000/health"),
        ("Connector reset", "POST /admin/connectors/{name}/reset"),
        ("DLQ reprocess", "POST /admin/dead-letter/reprocess/{record_id}"),
        ("Backup", "infrastructure/scripts/backup.sh"),
    ], ["Action", "Command / location"])

    h(doc, "14. Testing", 1)
    p(doc, "Pytest suites: test_normaliser, test_alert_engine, test_risk_engine, test_tenancy (scoping/delegation/derivation), test_benchmark, test_exports_ag, test_exports, test_new_endpoints. Run: cd backend && python -m pytest tests -v. CI runs them against Postgres + Redis services.")

    h(doc, "15. Demo accounts", 1)
    table(doc, [
        ("exec@example.com / pass123", "exec — whole estate"),
        ("soc@example.com / pass123", "soc"),
        ("appsec@example.com / pass123", "appsec"),
        ("dbsec@example.com / pass123", "dbsec"),
        ("compliance@example.com / pass123", "compliance"),
        ("transversal@example.com / pass123", "transversal-admin — whole estate"),
        ("provincesoc@example.com / pass123", "province-soc-lead — Gauteng"),
        ("admin@example.com / admin123", "admin — superadmin"),
    ], ["Credential", "Role / scope"])
    p(doc, "Login page supports a tenancy-scope override (province/department/branch).")

    h(doc, "16. Configuration reference", 1)
    p(doc, "Key env vars (backend/.env, see .env.example): DATABASE_URL, REDIS_URL, ENVIRONMENT, JWT_SECRET, AUTH0_DOMAIN/AUDIENCE, SEED_DEMO_USERS_ENABLED, BOOTSTRAP_ADMIN_EMAIL/PASSWORD, SMTP_*, TEAMS_WEBHOOK_URL, PAGERDUTY_ROUTING_KEY, DB_POOL_*, PROCESSING_*, DISPATCH_*, ANALYTICS_*. Non-dev environments are guarded to reject default dev secrets.")

    h(doc, "Appendix A — Related deliverables", 1)
    table(doc, [
        ("Stakeholder tenancy explainer", "docs/Tenancy_Stakeholder_Overview.docx"),
        ("Alert messaging design", "docs/SITA_Alert_Messaging.docx"),
        ("Database schema reference", "docs/database-schema.md"),
        ("Architecture tree + ER diagram", "tree.html"),
    ], ["Asset", "Location"])

    out = r"C:\sita-platform\docs\SITA_Platform_System_Documentation.docx"
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
