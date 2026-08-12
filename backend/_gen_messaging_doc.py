"""Generate SITA_Alert_Messaging.docx describing the alert-based messaging design."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor

SITA_BLUE = RGBColor(0x10, 0x82, 0xFF)
NAVY = RGBColor(0x23, 0x30, 0x42)
GREEN = RGBColor(0x00, 0xAF, 0x66)
RED = RGBColor(0xC0, 0x39, 0x2B)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else SITA_BLUE
    return h

def para(text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

# ---------------------------------------------------------------- title
title = doc.add_heading("SITA Security Intelligence Platform", level=0)
for run in title.runs:
    run.font.color.rgb = SITA_BLUE
sub = doc.add_paragraph()
r = sub.add_run("Alert-Based Messaging — Design & Operations Guide")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = NAVY
p = doc.add_paragraph()
p.add_run("Version 1.0  |  Date: 6 August 2026  |  Owner: SITA SOC Engineering")
p.runs[0].font.color.rgb = RGBColor(0x45, 0x55, 0x65)
doc.add_paragraph()

# ---------------------------------------------------------------- 1. overview
heading("1. Overview", 1)
para(
    "SITA's Security Intelligence Platform centralises security signals (AppScan findings, Imperva DAM/WAF "
    "events, API security, compliance scores, risk scores and connector health) and converts them into "
    "actionable alerts. The alert-based messaging layer is the delivery backbone: it enriches every alert "
    "with operational context, routes it to the correct notification channels, and records an auditable "
    "delivery trail for every attempt."
)
para(
    "This document describes the end-to-end alert messaging flow, the routing model, the enrichment and "
    "dispatch workers, the delivery adapters, the audit trail and the API/UI surface. It also records how "
    "the design was verified against seeded data."
)

# ---------------------------------------------------------------- 2. high level
heading("2. High-Level Flow", 1)
para("An alert moves through five stages from detection to audited delivery:", bold=False)
stages = [
    ("Detection", "AlertRuleEngine evaluates rules against findings, risk scores, API endpoints, compliance "
     "snapshots and connector health. When a rule condition fires, it builds a canonical alert record "
     "carrying id, rule_id, title, description, severity, source, target_id and the rule's declared channels."),
    ("Enrichment", "AlertEnricher attaches operational context (owner, team, tier, priority P1-P5, external "
     "id and a deep link into the SITA dashboard) using a config-driven application catalogue."),
    ("Publish", "The alert is written to the processing stream (sita:processing:alerts) on Redis as an "
     "envelope {id, type, source, ts, payload}."),
    ("Dispatch", "One or more DispatchWorker consumers read the stream, select the target channels for the "
     "alert, and invoke the relevant delivery adapters (email, Microsoft Teams, PagerDuty)."),
    ("Audit", "Every channel attempt is written to the warehouse.dispatch_log table and the alert's "
     "last_dispatched_at timestamp is stamped, producing a complete delivery history."),
]
table(["Stage", "What happens"], stages, widths=[1.1, 5.4])

# ---------------------------------------------------------------- 3. rule channels
heading("3. Rule-Declared Channels", 1)
para(
    "Each rule in backend/app/processing/alert_rules.yaml declares an explicit channels list. This is the "
    "policy source of truth: the dispatch worker honours exactly these channels and does not silently "
    "over- or under-deliver."
)
code('''critical_record       -> [email, teams, pagerduty]     (critical AppScan finding)
high_imperva_dam      -> [teams]                       (Imperva DAM event)
appscan_high          -> [teams]                       (high AppScan finding)
risk_critical         -> [email, teams, pagerduty]     (fused risk critical)
new_critical_cve      -> [email, teams, pagerduty]     (new critical CVE)
shadow_api            -> [email, teams]                (shadow API discovered)
compliance_drop       -> [email, teams]                (compliance score drop)
imperva_spike         -> [teams]                       (3-sigma violation spike)
connector_auth_failure-> [teams, pagerduty]            (connector down / degraded)''')
para(
    "Rule-declared channels give operators per-rule control: a noisy high-volume rule can be reduced to "
    "Teams-only (e.g. high_imperva_dam), while a business-critical condition can escalate to email, Teams "
    "and PagerDuty (e.g. risk_critical)."
)

# ---------------------------------------------------------------- 4. routing
heading("4. Routing & Severity Fallback", 1)
para(
    "The DispatchWorker.channels_for() method implements a two-tier routing model:"
)
bullet("Explicit channels win", "If the alert carries a non-empty channels list (from the rule), those channels are used exactly as declared.", )
bullet("Severity fallback", "Legacy or seed alerts without a channels list fall back to a severity-based default so they always reach someone.", )
table(
    ["Severity", "Fallback channels"],
    [
        ["critical", "email, teams, pagerduty"],
        ["high", "email, teams"],
        ["medium", "teams"],
        ["low", "teams"],
        ["info", "teams"],
    ],
    widths=[1.5, 3.0],
)
para(
    "Within each channel, adapter-level policy still applies: email is only dispatched for critical/high "
    "severity, PagerDuty only for critical severity. Teams accepts any severity. This layered control "
    "prevents an incorrect channel list from spamming low-priority noise into an on-call tool.",
    italic=True,
)

# ---------------------------------------------------------------- 5. enrichment
heading("5. Enrichment", 1)
para(
    "The AlertEnricher (backend/app/processing/alert_enricher.py) attaches operational context to every "
    "alert before it is published. Ownership is resolved from the application catalogue "
    "(app_catalog.yaml with built-in defaults for the five reference applications):"
)
table(
    ["Application", "Owner", "Team", "Tier"],
    [
        ["legacy-api", "AppSec Engineer", "AppSec", "P1"],
        ["payment-gateway", "Payment Ops", "Payments", "P1"],
        ["customer-portal", "Web Platform", "Digital", "P2"],
        ["document-svc", "Content Platform", "Digital", "P2"],
        ["internal-hr", "HR Systems", "Internal", "P3"],
    ],
    widths=[1.7, 1.5, 1.2, 0.8],
)
para("The enrichment payload delivered to every channel is:", bold=False)
code('''enriched_data = {
  "owner":         "Web Platform",
  "team":          "Digital",
  "tier":          "P2",
  "priority":      "P2",
  "external_id":   "<external or target id>",
  "dashboard_link": "/alerts/<alert-id>",
  "enriched_at":   "<iso timestamp>"
}''')
para(
    "This same enriched_data is persisted on the alert, returned by the API and rendered by the email "
    "adapter — so the recipient sees owner, team and a deep link, not just a raw message.",
)

# ---------------------------------------------------------------- 6. adapters
heading("6. Delivery Adapters", 1)
para("Three adapters implement the actual outbound delivery:", bold=False)
table(
    ["Adapter", "Module", "Behaviour"],
    [
        ["Email", "dispatch/email_adapter.py",
         "Renders an HTML template (severity, source, rule, owner, team, priority, description, deep link). "
         "Sends via SMTP with optional STARTTLS + login. Recipient configured to soc@sita.com."],
        ["Microsoft Teams", "dispatch/teams_adapter.py",
         "Posts a card to the configured Teams webhook URL. Accepts all severities."],
        ["PagerDuty", "dispatch/pagerduty_adapter.py",
         "Creates a critical-severity incident via the configured PagerDuty routing key."],
    ],
    widths=[1.3, 1.9, 3.6],
)
para(
    "Adapters return True/False. Dispatch executes channels concurrently, and every result (success or "
    "failure, including the exception message) is captured into a per-channel outcome.",
)

# ---------------------------------------------------------------- 7. audit
heading("7. Delivery Audit Trail", 1)
para(
    "Every dispatch attempt is recorded in the warehouse.dispatch_log table (model DispatchLog), giving the "
    "SOC full observability into whether notifications actually went out:"
)
table(
    ["Column", "Purpose"],
    [
        ["id", "Primary key (UUID)."],
        ["alert_id", "The alert the attempt belongs to (indexed)."],
        ["channel", "email | teams | pagerduty."],
        ["status", "sent | failed | skipped."],
        ["error", "Exception message or 'adapter returned False' when a send fails."],
        ["attempted_at", "UTC timestamp of the attempt."],
        ["created_at", "Row creation timestamp (server default)."],
    ],
    widths=[1.5, 5.0],
)
para(
    "In parallel, the dispatch entrypoint stamps the alert's last_dispatched_at column. The alerts table "
    "was extended with channels (JSON) and last_dispatched_at, so each alert carries both its routing "
    "policy and its most recent delivery time.",
)

# ---------------------------------------------------------------- 8. api
heading("8. API Surface", 1)
para("The alerts API exposes the messaging data to the frontend:", bold=False)
table(
    ["Endpoint", "What it returns"],
    [
        ["GET /api/v1/alerts", "Alert list including channels, enriched_data, dedup_count, last_dispatched_at "
         "and resolved_at."],
        ["GET /api/v1/alerts/{id}", "Full detail of a single alert, including enrichment."],
        ["GET /api/v1/alerts/{id}/dispatch", "The delivery audit trail — every channel attempt for the alert "
         "(channel, status, error, attempted_at)."],
    ],
    widths=[2.4, 4.1],
)

# ---------------------------------------------------------------- 9. ui
heading("9. Frontend (SOC Console)", 1)
para(
    "The SOC dashboard surfaces the messaging data so analysts can see ownership and delivery at a glance:"
)
bullet("A new Owner / Team column in the alert queue, populated from enriched_data.", "")
bullet("Alert detail views show priority, tier, external id and a deep link into the console.", "")
bullet("The Alert and DispatchLogEntry types in frontend/src/types.ts mirror the API schema.", "")

# ---------------------------------------------------------------- 10. config
heading("10. Configuration & Operations", 1)
para("Delivery depends on the following settings (backend/app/config.py):", bold=False)
table(
    ["Setting", "Purpose", "Current"],
    [
        ["smtp_host / smtp_port", "SMTP server for email delivery", "localhost:587"],
        ["smtp_user / smtp_password", "SMTP credentials (STARTTLS login when set)", "empty"],
        ["teams_webhook_url", "Teams Incoming Webhook URL", "empty"],
        ["pagerduty_routing_key", "PagerDuty Events routing key", "empty"],
        ["dispatch_consumers / dispatch_max_workers", "Number of dispatch consumers / threads", "configurable"],
    ],
    widths=[2.1, 3.2, 1.5],
)
para(
    "With the SMTP credentials, Teams webhook and PagerDuty key unset, adapters deliberately return False "
    "and the attempt is recorded as failed in the audit trail. This is by design: the platform records "
    "every attempt, and failures are never silently swallowed.",
)

# ---------------------------------------------------------------- 11. seed test
heading("11. Verification Against Seeded Data", 1)
para(
    "The messaging layer was verified end-to-end by driving seeded warehouse alerts through the "
    "DispatchWorker and recording the audit trail (backend/_test_dispatch.py):"
)
bullet("12 seeded alerts dispatched through DispatchWorker with their rule-declared channels.", "")
bullet("31 dispatch_log rows written (email x12, teams x12, pagerduty x7) and last_dispatched_at stamped on all 12 alerts.", "")
bullet("Per-channel statuses recorded (currently failed because no SMTP/webhook/PD credentials are configured).", "")
bullet("Routing assertions passed: a critical alert declaring [teams] delivers only to Teams; legacy alerts fall back to severity defaults.", "")
bullet("Email template renders owner, team, priority and dashboard link from enriched_data.", "")
para(
    "The full stream path (AlertRuleEngine → Redis sita:processing:alerts → dispatch consumer → audit) "
    "is the production flow; the seeded test exercises the same worker and audit logic directly against "
    "warehouse rows so it is runnable without Redis running.",
)

# ---------------------------------------------------------------- 12. summary
heading("12. Summary", 1)
para(
    "The alert-based messaging architecture gives SITA a single, auditable notification backbone: "
    "rule-declared channels provide precise per-rule routing, a severity fallback guarantees delivery for "
    "legacy alerts, enrichment carries ownership and deep links to every channel, and the dispatch_log "
    "audit trail provides full accountability for every notification attempt — sent or failed.",
)

doc.save(r"C:\sita-platform\docs\SITA_Alert_Messaging.docx")
print("Saved C:\\sita-platform\\docs\\SITA_Alert_Messaging.docx")
